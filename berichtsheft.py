from docx import Document
from datetime import datetime, timedelta
import os

#//----------------------------------Ausfüllen----------------------------------\\

template_path = "ausbildungsnachweis.docx"
output_folder = "Lehrjahr"

name = "Name"
schul_jahr = "Lehrjahr"
abteilung = "Ausbildung"
tätigkeit = "!AUSFÜLLEN!"

start_date = datetime(2024, 9, 2)
end_date = datetime(2025, 7, 30)

ferienzeiten = [
    (datetime(2024, 10, 28), datetime(2024, 11, 1)),  # Herbstferien
    (datetime(2024, 12, 23), datetime(2025, 1, 5)),   # Weihnachtsferien
    (datetime(2025, 2, 17), datetime(2025, 2, 21)),   # Fasnachtsferien
    (datetime(2025, 4, 7), datetime(2025, 4, 18)),    # Osterferien
    (datetime(2025, 5, 26), datetime(2025, 6, 6)),    # Pfingstferien
    (datetime(2025, 7, 28), datetime(2025, 9, 11)),   # Sommerferien
]

#//----------------------------------Code----------------------------------\\

def is_in_ferien(current_date):
    """Überprüft, ob das aktuelle Datum innerhalb der Ferien liegt"""
    for start, end in ferienzeiten:
        if start <= current_date <= end:
            return True
    return False

def fill_berichtsheft(template_path, output_folder):
    # Lade das Word-Dokument
    doc = Document(template_path)

    # Sicherstellen, dass der Zielordner existiert, andernfalls erstellen
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Schleife für jede Woche, bis zum Enddatum (30. Juli 2025)
    current_date = start_date
    week_num = 1

    while current_date <= end_date:
        # Berechne das Start- und Enddatum für die Woche (Montag bis Freitag)
        week_start_date = current_date
        week_end_date = week_start_date + timedelta(days=4)  # Eine Woche geht bis Freitag

        # Fülle die erste Tabelle mit den entsprechenden Daten
        table = doc.tables[0]  # Erste Tabelle im Dokument
        table.cell(0, 1).text = name  # Name des/der Auszubildenden
        table.cell(1, 1).text = schul_jahr  # Ausbildungsjahr
        table.cell(1, 4).text = abteilung  # Abteilung
        table.cell(2, 1).text = week_start_date.strftime("%d.%m.%Y")  # Ausbildungswoche vom
        table.cell(2, 3).text = week_end_date.strftime("%d.%m.%Y")  # bis

        # Wenn Ferienzeit ist, setze die Stunden auf 40, Schulstunden auf 0 und Schulungen/Fächer auf ---
        if is_in_ferien(week_start_date):
            stunden = 40
            schulstunden = 0
            schulungen_faecher = "---"
        else:
            # Wechsel zwischen 11 und 13 Stunden je Woche
            if week_num % 2 == 1:  # Woche 1, 3, 5, ... (ungerade Wochen)
                stunden = 27
                schulstunden = 11
                schulungen_faecher = "WK, GK, IT-LF-3 (Netzwerk), IT-LF-5 (Programmieren), IT-LF-1 (BWL), Religion, Deutsch, IT-LF-2 (Hardware)"
            else:  # Woche 2, 4, 6, ... (gerade Wochen)
                stunden = 27
                schulstunden = 13
                schulungen_faecher = "WK, GK, Deutsch, Englisch, IT-LF-1 (BWL), IT-LF-5 (Programmieren), Religion, IT-LF-3 (Netzwerk), IT-LF-5 (Datenbanken)"

        # Fülle die Tabellen für betriebliche Tätigkeiten und Schulungen aus
        if len(doc.tables) > 1:
            table_betrieblich = doc.tables[1]  # Tabelle für Betriebliche Tätigkeiten

            # Hier trägst du die Stunden und Tätigkeiten für jede Woche ein
            table_betrieblich.cell(1, 0).text = tätigkeit  # Tätigkeiten/Schulfächer
            table_betrieblich.cell(1, 1).text = f"{stunden}"  # Stunden
            table_betrieblich.cell(3, 0).text = "---"
            table_betrieblich.cell(3, 1).text = "0"
            table_betrieblich.cell(5, 0).text = schulungen_faecher
            table_betrieblich.cell(5, 1).text = f"{schulstunden}"

        # Erstellen des Dateipfads für jede Woche, um die Datei im gewünschten Ordner zu speichern
        week_output_path = os.path.join(output_folder, f"Woche_{week_num}.docx")
        doc.save(week_output_path)
        print(f"Das Berichtsheft für Woche {week_num} wurde erfolgreich gespeichert unter: {week_output_path}")

        # Zur nächsten Woche weitergehen
        current_date = week_end_date + timedelta(days=3)  # Gehe zum nächsten Montag
        week_num += 1

if __name__ == "__main__":
    fill_berichtsheft(template_path, output_folder)
